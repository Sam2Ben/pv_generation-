import os
import io
import base64
import streamlit as st
from PIL import Image
import tempfile
from pydub import AudioSegment
import subprocess
import time
import random
from google import generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import requests
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import json
from bs4 import BeautifulSoup
import av
from streamlit_webrtc import webrtc_streamer, WebRtcMode, AudioProcessorBase
from st_audiorec import st_audiorec
import tempfile
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import datetime
import re
from datetime import datetime
from docx.oxml.ns import qn
import concurrent.futures


# Configuration des APIs
def configure_apis():
    # Configure Google Gemini
    google_api_key = st.secrets.get("GOOGLE_API_KEY")
    if google_api_key:
        genai.configure(api_key=google_api_key)
    else:
        st.error("❌ Clé API Google non trouvée!")
        st.stop()

def convert_to_mp3(input_path, output_path):
    """Convertit n'importe quel format audio en MP3"""
    try:
        import shutil
        if not shutil.which("ffmpeg"):
            st.error("ffmpeg non trouvé. Veuillez installer ffmpeg.")
            return False
        audio = AudioSegment.from_file(input_path)
        audio.export(output_path, format="mp3")
        return True
    except Exception as e:
        st.error(f"Erreur de conversion audio : {e}")
        return False

def extract_file_id_from_url(url):
    """Extrait l'ID du fichier depuis une URL Google Drive"""
    patterns = [
        r"https://drive\.google\.com/file/d/([a-zA-Z0-9_-]+)",
        r"https://drive\.google\.com/open\?id=([a-zA-Z0-9_-]+)",
        r"https://drive\.google\.com/uc\?id=([a-zA-Z0-9_-]+)",
        r"id=([a-zA-Z0-9_-]+)"
    ]
    
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def verify_video_file(file_path):
    """Vérifie si le fichier vidéo est valide"""
    try:
        if not os.path.exists(file_path):
            st.error(f"❌ Le fichier {file_path} n'existe pas")
            return False
            
        file_size = os.path.getsize(file_path)
        if file_size < 10000:  # Moins de 10KB est suspect
            st.error("❌ Le fichier est trop petit pour être une vidéo valide")
            return False
            
        st.info(f"📊 Taille du fichier : {file_size/1024/1024:.1f} MB")
        
        # Vérification du format avec ffprobe
        probe_command = [
            "ffprobe",
            "-v", "error",
            "-show_format",
            "-show_streams",
            file_path
        ]
        
        result = subprocess.run(probe_command, capture_output=True, text=True)
        
        if result.returncode != 0:
            st.error(f"❌ Format vidéo non valide: {result.stderr}")
            return False
            
        st.success("✅ Format vidéo validé")
        return True
            
    except Exception as e:
        st.error(f"❌ Erreur lors de la vérification: {str(e)}")
        return False

def convert_vro_to_mp4(input_path, output_path):
    """Convertit un fichier VRO en MP4"""
    try:
        st.info("🔄 Conversion du fichier VRO en MP4...")
        
        # Commande de conversion optimisée pour les fichiers VRO
        convert_command = [
            "ffmpeg",
            "-y",  # Écraser le fichier de sortie si existant
            "-fflags", "+genpts",  # Générer les timestamps
            "-i", input_path,
            "-c:v", "libx264",  # Codec vidéo
            "-preset", "ultrafast",  # Conversion rapide
            "-crf", "23",  # Qualité raisonnable
            "-c:a", "aac",  # Codec audio
            "-strict", "experimental",
            "-b:a", "192k",  # Bitrate audio
            "-movflags", "+faststart",  # Optimisation pour la lecture web
            output_path
        ]
        
        # Exécuter la conversion
        result = subprocess.run(convert_command, capture_output=True, text=True)
        
        if result.returncode != 0:
            st.error(f"❌ Erreur lors de la conversion VRO: {result.stderr}")
            return False
            
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            st.success("✅ Conversion VRO → MP4 réussie")
            return True
        else:
            st.error("❌ Fichier MP4 non créé ou vide")
            return False
            
    except Exception as e:
        st.error(f"❌ Erreur lors de la conversion: {str(e)}")
        return False

def extract_audio_from_video(input_video_path, output_audio_path):
    """Extrait l'audio d'une vidéo"""
    try:
        # Vérifier si le fichier existe
        if not os.path.exists(input_video_path):
            st.error("❌ Le fichier vidéo n'existe pas")
            return False
            
        # Vérifier si le fichier est vide
        if os.path.getsize(input_video_path) == 0:
            st.error("❌ Le fichier vidéo est vide")
            return False
            
        # Si c'est un fichier VRO, on le convertit d'abord en MP4
        if input_video_path.lower().endswith('.vro'):
            st.info("🔄 Conversion du fichier VRO en MP4...")
            temp_mp4 = input_video_path + '.mp4'
            try:
                # Commande de conversion VRO vers MP4
                convert_command = [
                    'ffmpeg',
                    '-i', input_video_path,
                    '-c:v', 'libx264',
                    '-preset', 'ultrafast',
                    '-c:a', 'aac',
                    '-strict', 'experimental',
                    '-write_xing', '0',
                    '-y',
                    temp_mp4
                ]
                
                result = subprocess.run(convert_command, capture_output=True, text=True)
                if result.returncode != 0:
                    st.error(f"❌ Erreur lors de la conversion VRO vers MP4: {result.stderr}")
                    return False
                    
                input_video_path = temp_mp4
                st.success("✅ Conversion VRO vers MP4 réussie")
                
            except Exception as e:
                st.error(f"❌ Erreur lors de la conversion VRO vers MP4: {str(e)}")
                return False
                
        # Extraction de l'audio
        st.info("🎵 Extraction de l'audio...")
        try:
            # Commande d'extraction audio
            extract_command = [
                'ffmpeg',
                '-i', input_video_path,
                '-vn',
                '-acodec', 'libmp3lame',
                '-ar', '44100',
                '-ab', '192k',
                '-y',
                output_audio_path
            ]
            
            result = subprocess.run(extract_command, capture_output=True, text=True)
            if result.returncode != 0:
                st.error(f"❌ Erreur lors de l'extraction audio: {result.stderr}")
                return False
                
            # Vérifier si le fichier audio a été créé et n'est pas vide
            if not os.path.exists(output_audio_path) or os.path.getsize(output_audio_path) == 0:
                st.error("❌ Le fichier audio n'a pas été créé ou est vide")
                return False
                
            st.success("✅ Extraction audio réussie")
            return True
            
        except Exception as e:
            st.error(f"❌ Erreur lors de l'extraction audio: {str(e)}")
            return False
            
    except Exception as e:
        st.error(f"❌ Erreur inattendue lors du traitement: {str(e)}")
        return False
    finally:
        # Nettoyage des fichiers temporaires
        try:
            temp_mp4 = input_video_path + '.mp4'
            if os.path.exists(temp_mp4):
                os.remove(temp_mp4)
        except Exception as e:
            st.warning(f"⚠️ Erreur lors du nettoyage des fichiers temporaires: {str(e)}")

def segment_audio(audio_path, segment_length_ms=120000):
    """Divise un gros fichier audio en segments sans tout charger en RAM (version ultra-optimisée)"""
    try:
        import math
        import subprocess
        
        # Utiliser ffprobe pour obtenir la durée
        result = subprocess.run([
            'ffprobe', '-v', 'error', '-show_entries',
            'format=duration', '-of',
            'default=noprint_wrappers=1:nokey=1', audio_path
        ], stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
        
        total_duration = float(result.stdout)
        segment_length_sec = segment_length_ms / 1000
        
        num_segments = math.ceil(total_duration / segment_length_sec)
        segment_paths = []  # Stocker uniquement les chemins

        temp_dir = tempfile.gettempdir()

        for i in range(num_segments):
            start_time = i * segment_length_sec
            temp_segment_path = os.path.join(temp_dir, f"segment_{i+1}.mp3")
            
            extract_cmd = [
                "ffmpeg",
                "-y",
                "-i", audio_path,
                "-ss", str(start_time),
                "-t", str(segment_length_sec),
                "-c", "copy",
                temp_segment_path
            ]
            subprocess.run(extract_cmd, stdout=subprocess.DEVNULL, stderr=subprocess.STDOUT)
            
            if os.path.exists(temp_segment_path):
                segment_paths.append(temp_segment_path)

        return segment_paths
        
    except Exception as e:
        st.error(f"❌ Erreur lors de la segmentation audio (stream) : {str(e)}")
        return []

        
    except Exception as e:
        st.error(f"❌ Erreur lors de la segmentation audio (stream) : {str(e)}")
        return []


def process_segment_batch(segments, start_idx, batch_size, total_segments, progress_bar, status_text):
    """Traite un lot de segments audio (optimisé sans chargement mémoire) avec timeout par segment et feedback UX amélioré."""
    batch_transcript = []
    SEGMENT_TIMEOUT = 30  # secondes
    start_time = time.time()
    for i in range(start_idx, min(start_idx + batch_size, total_segments)):
        segment_path = segments[i]  # Maintenant segments contient des chemins de fichier
        segment_number = i + 1
        try:
            status_text.text(f"🎯 Traitement du segment {segment_number}/{total_segments}")
            with open(segment_path, "rb") as f:
                audio_bytes = f.read()
            model = genai.GenerativeModel('gemini-2.0-flash')
            def call_gemini():
                return model.generate_content([
                    "Transcrivez ce segment audio mot pour mot en français.",
                    {"mime_type": "audio/mp3", "data": audio_bytes}
                ])
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(call_gemini)
                try:
                    response = future.result(timeout=SEGMENT_TIMEOUT)
                    if response.text:
                        batch_transcript.append(response.text)
                        progress_bar.progress((i + 1)/total_segments)
                    else:
                        batch_transcript.append(f"[Segment {segment_number} non transcrit]")
                except concurrent.futures.TimeoutError:
                    st.warning(f"⏰ Timeout sur le segment {segment_number} (>{SEGMENT_TIMEOUT}s)")
                    batch_transcript.append(f"[Segment {segment_number} timeout]")
            os.remove(segment_path)
        except Exception as e:
            st.warning(f"⚠️ Erreur sur le segment {segment_number}: {str(e)}")
            batch_transcript.append(f"[Segment {segment_number} non transcrit]")
        elapsed = time.time() - start_time
        if elapsed > 120:
            st.warning("⏳ Le traitement audio prend plus de 2 minutes. Merci de patienter ou essayez avec un fichier plus court.")
        time.sleep(random.uniform(1, 2))  # Attente pour respecter quotas API
    status_text.text("Traitement du lot terminé.")
    return batch_transcript


def transcribe_video(video_file):
    """Transcrit une vidéo en texte sans charger tout en RAM."""
    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            # Créer un chemin temporaire avec bonne extension
            if hasattr(video_file, 'name'):
                ext = os.path.splitext(video_file.name)[1].lower()
                video_temp_path = os.path.join(temp_dir, f"uploaded_video{ext}")
                st.info(f"📝 Sauvegarde du fichier temporaire: {video_file.name}")
            else:
                ext = '.mp4'
                video_temp_path = os.path.join(temp_dir, "uploaded_video.mp4")
                st.info("📝 Sauvegarde d'un fichier vidéo sans nom")

            # Maintenant, on enregistre directement l'objet téléchargé
            with open(video_temp_path, 'wb') as out_file:
                for chunk in iter(lambda: video_file.read(1024 * 1024), b''):
                    out_file.write(chunk)

            st.success("✅ Vidéo sauvegardée sur disque temporaire")

            # On vérifie la taille
            video_size = os.path.getsize(video_temp_path)
            st.info(f"📊 Taille du fichier vidéo: {video_size/1024/1024:.2f} MB")

            # Vérifier la validité
            if not verify_video_file(video_temp_path):
                return ""

            # On extrait maintenant l'audio
            audio_path = os.path.join(temp_dir, "output_audio.mp3")
            st.info("🎵 Extraction de l'audio...")
            if not extract_audio_from_video(video_temp_path, audio_path):
                return ""

            # Segmentation
            st.info("🔄 Segmentation de l'audio...")
            segments = segment_audio(audio_path)
            if not segments:
                st.error("❌ Échec de la segmentation audio")
                return ""

            st.success(f"✅ Audio segmenté en {len(segments)} parties")
            
            # Traitement par batch
            progress_bar = st.progress(0)
            status_text = st.empty()

            full_transcript = []
            BATCH_SIZE = 10

            for batch_start in range(0, len(segments), BATCH_SIZE):
                batch_results = process_segment_batch(
                    segments, batch_start, BATCH_SIZE, len(segments),
                    progress_bar, status_text
                )
                full_transcript.extend(batch_results)

            if not full_transcript:
                st.warning("⚠️ Aucun texte n'a été transcrit")
                return ""
                
            st.success("✅ Transcription terminée avec succès")
            return "\n".join(full_transcript)

        except Exception as e:
            st.error(f"❌ Erreur lors de la transcription: {str(e)}")
            return ""


def process_handwritten_image(image_bytes):
    """Extrait le texte d'une image manuscrite avec mécanisme de retry"""
    @retry_with_backoff
    def transcribe_image():
        try:
            image_base64 = base64.b64encode(image_bytes).decode('utf-8')
            
            model = genai.GenerativeModel('gemini-2.0-flash')
            
            prompt = """Transcris précisément le texte manuscrit dans cette image.
            INSTRUCTIONS :
            1. Retourne uniquement le texte, sans commentaires
            2. Préserve la mise en forme (retours à la ligne, espacements)
            3. Conserve la ponctuation exacte
            4. Maintiens les nombres et symboles tels quels
            5. Respecte les majuscules et minuscules"""
            
            response = model.generate_content([
                prompt,
                {"mime_type": "image/jpeg", "data": image_base64}
            ])
            
            if response.text:
                return response.text.strip()
            else:
                raise Exception("Aucun texte détecté dans l'image.")
                
        except Exception as e:
            st.warning(f"⚠️ Tentative de transcription échouée : {str(e)}")
            raise e

    try:
        # Premier essai
        result = transcribe_image()
        if result:
            return result
            
        # Si le résultat est vide, on attend et on réessaie
        time.sleep(2)  # Attente de 2 secondes
        st.info("🔄 Nouvelle tentative de transcription...")
        
        # Deuxième essai avec un prompt plus détaillé
        prompt_retry = """Analyse et transcris TOUT le texte manuscrit visible dans cette image.
        IMPORTANT :
        - Examine l'image en détail, pixel par pixel
        - Transcris absolument tout le texte visible
        - N'oublie aucun détail, même les petites annotations
        - Conserve la structure exacte du texte
        - Inclus les numéros, symboles et caractères spéciaux"""
        
        model = genai.GenerativeModel('gemini-2.0-flash')
        image_base64 = base64.b64encode(image_bytes).decode('utf-8')
        
        response = model.generate_content([
            prompt_retry,
            {"mime_type": "image/jpeg", "data": image_base64}
        ])
        
        if response.text:
            return response.text.strip()
        else:
            st.warning("⚠️ Aucun texte n'a été détecté dans l'image après plusieurs tentatives.")
            return ""
            
    except Exception as e:
        st.error(f"❌ Erreur lors de la reconnaissance du texte : {str(e)}")
        return ""

def retry_with_backoff(func, max_retries=5, initial_delay=1):
    """Fonction utilitaire pour réessayer une opération avec un délai exponentiel"""
    def wrapper(*args, **kwargs):
        delay = initial_delay
        last_exception = None
        
        for attempt in range(max_retries):
            try:
                return func(*args, **kwargs)
            except Exception as e:
                last_exception = e
                error_code = str(e)
                # Réessayer sur les erreurs de quota (429) ou d'annulation client (499)
                if "429" in error_code or "499" in error_code: 
                    st.warning(f"⚠️ Erreur API ({error_code}), nouvelle tentative {attempt + 1}/{max_retries} dans {delay} secondes...")
                    time.sleep(delay)
                    delay *= 2  # Backoff exponentiel
                else:
                    # Pour les autres exceptions, ne pas réessayer
                    raise e
        
        st.error(f"❌ Échec après {max_retries} tentatives : {str(last_exception)}")
        # Retourner None ou une valeur indiquant l'échec si toutes les tentatives échouent
        return None
    
    return wrapper

def process_pdf(pdf_file):
    """Extrait le contenu détaillé et les acronymes d'un PDF en un seul appel."""
    try:
        pdf_bytes = pdf_file.read()
        pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
        
        model = genai.GenerativeModel('gemini-2.0-flash')
        
        prompt = """Analyse ce document PDF de manière EXHAUSTIVE et DÉTAILLÉE.
        
        INSTRUCTIONS SPÉCIFIQUES :
        
        1. EXTRACTION COMPLÈTE DU CONTENU :
           - Extraire TOUS les textes, exactement comme ils apparaissent.
           - Conserver TOUS les chiffres, statistiques, données numériques avec leurs unités.
           - Maintenir TOUS les tableaux avec leurs données complètes.
           - Décrire TOUS les graphiques avec leurs valeurs précises.
           - Capturer TOUTES les notes de bas de page et références.
           - Respecter la structure (sections, titres, listes).
           - NE PAS résumer ou synthétiser le corps du texte.
           
        2. EXTRACTION DES ACRONYMES :
           - Identifier TOUS les acronymes présents dans le document.
           - Si l'acronyme est défini explicitement dans le texte, utiliser cette définition EXACTE.
           - Si l'acronyme n'est PAS défini dans le texte, rechercher sa définition officielle connue dans des sources fiables.
           - Lister les acronymes et leurs définitions SÉPARÉMENT à la fin.
        
        3. FORMAT DE SORTIE ATTENDU :
           - D'abord, le contenu complet et détaillé du document, en respectant sa structure.
           - Ensuite, une ligne de séparation claire comme : '--- ACRONYMES ---'.
           - Enfin, la liste des acronymes, un par ligne, au format : 'ACRONYME: Définition complète'.
           
        IMPORTANT : Assure-toi de bien séparer le contenu principal de la liste des acronymes avec '--- ACRONYMES ---'."""
        
        @retry_with_backoff
        def analyze_pdf_and_extract_acronyms():
            response = model.generate_content([
                {
                    "role": "user",
                    "parts": [
                        prompt,
                        {"mime_type": "application/pdf", "data": pdf_base64}
                    ]
                }
            ])
            return response.text if response.text else ""
        
        full_result = analyze_pdf_and_extract_acronyms()
        
        if not full_result:
            st.warning(f"⚠️ Aucun contenu extrait du PDF: {pdf_file.name}")
            return {"summary": "", "acronyms": {}}
            
        # Séparer le contenu et les acronymes
        separator = "--- ACRONYMES ---"
        if separator in full_result:
            summary_part, acronym_part = full_result.split(separator, 1)
            summary = summary_part.strip()
            
            # Parser les acronymes
            acronyms = {}
            lines = acronym_part.strip().split('\n')
            for line in lines:
                if ':' in line:
                    acronym, definition = line.split(':', 1)
                    acronym = acronym.strip().upper()
                    definition = definition.strip()
                    if acronym and definition:
                        acronyms[acronym] = definition
            return {"summary": summary, "acronyms": acronyms}
        else:
            # Si le séparateur n'est pas trouvé, retourner tout comme résumé et pas d'acronymes
            st.warning(f"⚠️ Séparateur d'acronymes non trouvé dans l'analyse de {pdf_file.name}")
            return {"summary": full_result.strip(), "acronyms": {}}
            
    except Exception as e:
        st.error(f"❌ Erreur lors de l'analyse du PDF {pdf_file.name}: {str(e)}")
        return {"summary": f"[Erreur lors de l'analyse du PDF: {str(e)}]", "acronyms": {}}


def create_word_pv(pv_text, meeting_info):
    doc = Document()

    # === En-tête centré ===
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_text = (
        'TANGER MED PORT AUTHORITY S.A "TMPA"\n'
        'SOCIÉTÉ ANONYME À CONSEIL D\'ADMINISTRATION\n'
        'AU CAPITAL DE 1.704.000.000 DIRHAMS CONVERTIBLES\n'
        'SIÈGE SOCIAL : ZONE FRANCHE DE KSAR EL MAJAZ, OUED RMEL,\n'
        'COMMUNE ANJRA ROUTE DE FNIDEQ – TANGER\n'
        'RC N°45349 TANGER – ICE : 000053443000022'
    )
    header_para.text = header_text
    for run in header_para.runs:
        run.font.size = Pt(9)
        run.bold = True
        run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # === Titre centré ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROCÈS VERBAL DE LA RÉUNION DU CONSEIL D'ADMINISTRATION\n")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'

    # Date centrée
    try:
        date_str = meeting_info.get("date", "")
        date_obj = datetime.strptime(date_str, "%d/%m/%Y")
        formatted_date = date_obj.strftime("DU %d %B %Y").upper()
    except:
        formatted_date = f"DU {date_str}"
    date_p = doc.add_paragraph(formatted_date)
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].bold = True
    date_p.runs[0].font.size = Pt(12)

    doc.add_paragraph()

    # === Introduction formelle ===
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    try:
        year = date_obj.year % 100
        date_lettres = date_obj.strftime('%d %B').capitalize()
        heure = meeting_info.get('heure_debut', '')
        lieu = meeting_info.get('lieu', '')
    except:
        year = 'XX'
        date_lettres = '[jour mois écrit en lettres]'
        heure = '[Heure]'
        lieu = '[Lieu]'
    intro_text = f"L'An Deux Mille {year}, Le {date_lettres}, À {heure} heures."
    intro.add_run(intro_text)
    doc.add_paragraph()
    doc.add_paragraph(f"Les membres du Conseil d'Administration de Tanger Med Port Authority S.A, par abréviation, « TMPA » se sont réunis en Conseil d'Administration en présentiel {('au ' + lieu) if lieu else ''} sur convocation et sous la présidence de.")
    doc.add_paragraph()

    # === Participants ===
    # Extraire les participants du texte généré
    participants_section = ""
    if "PARTICIPANTS" in pv_text:
        start_idx = pv_text.find("PARTICIPANTS")
        end_idx = pv_text.find("ORDRE DU JOUR")
        if end_idx == -1:
            end_idx = len(pv_text)
        participants_section = pv_text[start_idx:end_idx].strip()

    # Ajouter les participants présents
    if "Présents" in participants_section:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("SONT PRÉSENTS OU REPRÉSENTÉS :")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        
        # Extraire les présents
        start_idx = participants_section.find("Présents")
        end_idx = participants_section.find("Absents")
        if end_idx == -1:
            end_idx = len(participants_section)
        presents_text = participants_section[start_idx:end_idx].strip()
        
        # Ajouter chaque participant
        for line in presents_text.split('\n'):
            if line.strip() and not line.startswith("Présents"):
                para = doc.add_paragraph(line.strip(), style='List Bullet')
                para.paragraph_format.left_indent = Pt(24)

    # Ajouter les absents
    if "Absents" in participants_section:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("EST ABSENT EXCUSÉ")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        
        # Extraire les absents
        start_idx = participants_section.find("Absents")
        end_idx = participants_section.find("Invités")
        if end_idx == -1:
            end_idx = len(participants_section)
        absents_text = participants_section[start_idx:end_idx].strip()
        
        # Ajouter chaque absent
        for line in absents_text.split('\n'):
            if line.strip() and not line.startswith("Absents"):
                para = doc.add_paragraph(line.strip(), style='List Bullet')
                para.paragraph_format.left_indent = Pt(24)

    # Ajouter les invités
    if "Invités" in participants_section:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("ASSISTENT ÉGALEMENT À LA RÉUNION :")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        
        # Extraire les invités
        start_idx = participants_section.find("Invités")
        invites_text = participants_section[start_idx:].strip()
        
        # Ajouter chaque invité
        for line in invites_text.split('\n'):
            if line.strip() and not line.startswith("Invités"):
                para = doc.add_paragraph(line.strip(), style='List Bullet')
                para.paragraph_format.left_indent = Pt(24)

    doc.add_paragraph()
    doc.add_paragraph("Une feuille de présence a été établie et signée conformément à la loi par les membres du Conseil d'Administration participant à la réunion, chacun tant en son nom personnel que comme mandataire dûment habilité selon un pouvoir spécial.")
    doc.add_paragraph()

    # === Ordre du jour ===
    doc.add_paragraph("ORDRE DU JOUR :", style='Normal').runs[0].bold = True
    
    # Extraire l'ordre du jour du texte généré
    if "ORDRE DU JOUR" in pv_text:
        start_idx = pv_text.find("ORDRE DU JOUR")
        end_idx = pv_text.find("DÉROULÉ DE LA RÉUNION")
        if end_idx == -1:
            end_idx = len(pv_text)
        ordre_du_jour_text = pv_text[start_idx:end_idx].strip()
        
        # Ajouter chaque point de l'ordre du jour
        for line in ordre_du_jour_text.split('\n'):
            if line.strip() and not line.startswith("ORDRE DU JOUR"):
                if line.strip().startswith(('1.', '2.', '3.', '4.', '5.')):
                    para = doc.add_paragraph(line.strip(), style='List Number')
                    para.paragraph_format.left_indent = Pt(24)

    doc.add_paragraph()

    # === Déroulé de la réunion ===
    if "DÉROULÉ DE LA RÉUNION" in pv_text:
        start_idx = pv_text.find("DÉROULÉ DE LA RÉUNION")
        end_idx = pv_text.find("CONCLUSION")
        if end_idx == -1:
            end_idx = len(pv_text)
        deroule_text = pv_text[start_idx:end_idx].strip()
        
        # Traiter chaque point
        current_point = None
        for line in deroule_text.split('\n'):
            line = line.strip()
            if not line or line.startswith("DÉROULÉ DE LA RÉUNION"):
                continue
                
            if line.startswith("Point"):
                if current_point:
                    doc.add_paragraph()
                current_point = doc.add_paragraph()
                run = current_point.add_run(line.split(":", 1)[1].strip().upper())
                run.bold = True
                run.underline = True
                run.font.color.rgb = RGBColor(0, 32, 96)
            elif line.startswith(("Décisions", "Discussions", "Résolutions")):
                para = doc.add_paragraph()
                run = para.add_run(line.split(":", 1)[0] + " :")
                run.bold = True
                content = line.split(":", 1)[1].strip()
                if content:
                    para.add_run(" " + content)
            else:
                if current_point:
                    doc.add_paragraph(line)

    # === Conclusion ===
    if "CONCLUSION" in pv_text:
        start_idx = pv_text.find("CONCLUSION")
        conclusion_text = pv_text[start_idx:].strip()
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("CONCLUSION")
        run.bold = True
        run.font.size = Pt(12)
        
        for line in conclusion_text.split('\n'):
            if line.strip() and not line.startswith("CONCLUSION"):
                doc.add_paragraph(line.strip())

    # === Pied de page ===
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"PV_CA_TMPA_{meeting_info.get('date', '').replace('/', '_')}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_meeting_minutes(video_transcript, handwritten_text, pdf_summary, meeting_info, audio_transcript=None):
    """Génère un PV de réunion structuré avec un niveau de détail élevé et des données précises"""
    try:
        # Formater les sources d'information de manière plus structurée
        combined_text = ""
        
        # 1. Ajouter la transcription vidéo si disponible
        if video_transcript and video_transcript.strip():
            combined_text += "[TRANSCRIPTION VIDÉO]\n"
            combined_text += video_transcript.strip() + "\n\n"
        
        # 2. Ajouter les notes manuscrites si disponibles
        if handwritten_text and handwritten_text.strip():
            combined_text += "[NOTES MANUSCRITES]\n"
            combined_text += handwritten_text.strip() + "\n\n"
        
        # 3. Ajouter le contenu des documents PDF si disponible
        if pdf_summary and pdf_summary.strip():
            combined_text += "[DOCUMENTS PDF]\n"
            combined_text += pdf_summary.strip() + "\n\n"

        # 4. Ajouter la transcription audio si disponible   
        if audio_transcript and audio_transcript.strip():
            combined_text += "[ENREGISTREMENT AUDIO]\n"
            combined_text += audio_transcript.strip() + "\n\n"  

        if not combined_text.strip():
            return "Aucun contenu disponible pour générer le PV."

        model = genai.GenerativeModel('gemini-2.0-flash')
        
        prompt = f"""Analyse les sources d'information suivantes et génère un procès-verbal structuré au format suivant :

1. PARTICIPANTS :
   - Présents : [Liste des participants présents]
   - Absents excusés : [Liste des absents excusés]
   - Invités : [Liste des invités]

2. ORDRE DU JOUR :
   [Liste numérotée des points à l'ordre du jour]

3. DÉROULÉ DE LA RÉUNION :
   Pour chaque point de l'ordre du jour :
   - Titre du point
   - Décisions prises
   - Discussions importantes
   - Résolutions adoptées

4. CONCLUSION :
   - Résumé des décisions principales
   - Prochaines étapes
   - Date de la prochaine réunion si mentionnée

Sources d'information :
{combined_text}

Instructions :
1. Utilise UNIQUEMENT les informations présentes dans les sources
2. Respecte la structure demandée
3. Sois précis et professionnel
4. Inclus toutes les décisions et discussions importantes
5. Mentionne les votes et résolutions si présents dans les sources"""

        @retry_with_backoff
        def generate_content():
            response = model.generate_content([
                {
                    "role": "user",
                    "parts": [prompt]
                }
            ])
            return response.text if response.text else ""

        result = generate_content()
        
        if result:
            # Nettoyage et formatage du texte généré
            result = result.replace('**', '')
            result = result.replace('*', '')
            
            # Formater les titres de section
            for i in range(1, 10):
                result = result.replace(f'{i}.\n', f'{i}. ')
                result = result.replace(f'\n{i}. \n', f'\n{i}. ')
                result = result.replace(f'\n{i}.\n', f'\n{i}. ')
                result = result.replace(f'\n{i}.', f'\n\n{i}.')
            
            # Formater l'ordre du jour
            if "ORDRE DU JOUR" in result:
                ordre_index = result.index("ORDRE DU JOUR")
                result = result[:ordre_index] + "\n\nORDRE DU JOUR :\n" + result[ordre_index + 13:]
            
            # Formater les listes à puces
            result = result.replace('• ', '\n• ')
            
            # Assurer des sauts de ligne appropriés
            result = result.replace('\n\n\n', '\n\n')
            
            # Vérification finale pour les numéros isolés
            result = re.sub(r'\n(\d+\.)\s*\n', r'\n\1 ', result)
            
            return result.strip()
        else:
            st.warning("⚠️ Aucun contenu n'a été généré pour le PV.")
            return ""
                
    except Exception as e:
        st.error(f"❌ Erreur lors de la génération du PV : {str(e)}")
        return ""

def download_video_from_drive(video_url, output_path):
    """Télécharge une vidéo depuis Google Drive avec gestion des gros fichiers"""
    try:
        status_box = st.empty()
        progress_bar = st.empty()
        status_box.info("🔄 Initialisation du téléchargement...")
        
        # Extraire l'ID du fichier
        file_id = extract_file_id_from_url(video_url)
        if not file_id:
            st.error("❌ Format d'URL Google Drive non reconnu")
            return False

        status_box.info(f"📝 ID du fichier extrait : {file_id}")

        # Configuration de la session avec des headers complets
        session = requests.Session()
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'fr,fr-FR;q=0.8,en-US;q=0.5,en;q=0.3',
            'Accept-Encoding': 'gzip, deflate, br',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1'
        }

        # Utiliser l'URL de téléchargement direct avec usercontent
        download_url = f'https://drive.usercontent.google.com/download?id={file_id}&export=download&authuser=0&confirm=t'
        status_box.info(f"🔍 Téléchargement en cours...")
        
        response = session.get(download_url, headers=headers, stream=True, timeout=30)
        content_type = response.headers.get('Content-Type', '').lower()
        if 'text/html' in content_type:
            status_box.warning("⚠️ Redirection vers la page de confirmation détectée. Tentative alternative...")
            # Essayer l'URL alternative pour les gros fichiers
            download_url = f'https://drive.usercontent.google.com/download?id={file_id}&export=download&authuser=0&confirm=t&uuid=123&at=123'
            response = session.get(download_url, headers=headers, stream=True, timeout=30)
            content_type = response.headers.get('Content-Type', '').lower()
            if 'text/html' in content_type:
                st.error("❌ Impossible d'accéder au fichier. Assurez-vous que :\n1. Le fichier est partagé avec 'Tout le monde avec le lien'\n2. Vous avez les droits 'Lecteur' sur le fichier\n3. Le fichier n'est pas dans la corbeille")
                return False

        # Utiliser un nom de fichier temporaire unique
        import tempfile
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, f"download_{file_id}_{int(time.time())}.tmp")
        
        try:
            chunk_size = 500 * 1024 * 1024  # 500MB chunks
            downloaded_size = 0
            expected_size = None
            if 'content-length' in response.headers:
                expected_size = int(response.headers['content-length'])
                status_box.info(f"📦 Taille totale du fichier : {expected_size/1024/1024:.1f} MB")
            else:
                status_box.info("📦 Taille totale du fichier inconnue")

            with open(temp_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=chunk_size):
                    if chunk:
                        f.write(chunk)
                        downloaded_size += len(chunk)
                        # Afficher la progression
                        if expected_size:
                            current_percent = (downloaded_size/expected_size)
                            status_box.info(f"📥 Téléchargé : {downloaded_size/1024/1024:.1f} MB / {expected_size/1024/1024:.1f} MB ({current_percent*100:.1f}%)")
                            progress_bar.progress(min(1.0, current_percent))
                        else:
                            status_box.info(f"📥 Téléchargé : {downloaded_size/1024/1024:.1f} MB")

            # Vérifier le fichier téléchargé
            if os.path.exists(temp_path):
                file_size = os.path.getsize(temp_path)
                if file_size < 10000:  # Moins de 10KB
                    st.error("❌ Fichier téléchargé invalide ou trop petit")
                    try:
                        os.remove(temp_path)
                    except:
                        pass
                    return False
                # Vérifier les premiers octets pour s'assurer que c'est un fichier VRO
                with open(temp_path, 'rb') as f:
                    header = f.read(8)
                    if not header.startswith(b'DVD') and not header.startswith(b'\x00\x00\x01\xBA'):
                        st.error("❌ Le fichier téléchargé n'est pas un fichier VRO valide")
                        try:
                            os.remove(temp_path)
                        except:
                            pass
                        return False
                # Renommer le fichier temporaire
                try:
                    if os.path.exists(output_path):
                        os.remove(output_path)
                    os.rename(temp_path, output_path)
                except Exception as e:
                    st.error(f"❌ Erreur lors du déplacement du fichier : {str(e)}")
                    # Essayer de copier le fichier à la place
                    import shutil
                    try:
                        shutil.copy2(temp_path, output_path)
                        os.remove(temp_path)
                    except Exception as e2:
                        st.error(f"❌ Échec de la copie du fichier : {str(e2)}")
                        return False
                st.success(f"✅ Téléchargement réussi - Taille : {file_size/1024/1024:.1f} MB")
                return True
            else:
                st.error("❌ Échec de l'écriture du fichier")
                return False
        except Exception as e:
            st.error(f"❌ Erreur pendant le téléchargement : {str(e)}")
            try:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            except:
                pass
            return False
    except Exception as e:
        st.error(f"❌ Erreur inattendue : {str(e)}")
        try:
            if os.path.exists(output_path):
                os.remove(output_path)
        except:
            pass
        return False


def record_audio_simple():
    st.subheader("🎤 Enregistrement vocal")
    wav_audio_data = st_audiorec()

    if wav_audio_data:
        st.success("✅ Enregistrement terminé !")
        
        # Utiliser tempfile pour créer un fichier temporaire pour l'audio enregistré
        try:
            st.info("Création du fichier temporaire...")
            # Utilisez NamedTemporaryFile avec delete=False pour que le fichier persiste après la fermeture
            with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_audio_file:
                temp_audio_path = temp_audio_file.name
                st.info(f"Fichier temporaire créé : {temp_audio_path}")
                
                st.info("Écriture des données audio dans le fichier temporaire...")
                temp_audio_file.write(wav_audio_data)
                st.info("Écriture terminée.")
            
            # Stocker le chemin du fichier temporaire dans la session
            st.session_state.audio_file_path = temp_audio_path
            st.session_state.audio_transcript = ""  # Reset transcript
            st.write(f"📂 Chemin local du fichier temporaire : `{temp_audio_path}`")

            # Afficher les options de lecture et téléchargement à partir du fichier temporaire
            # Il faut rouvrir le fichier pour le lire car il a été fermé par le 'with' statement
            st.audio(open(temp_audio_path, "rb").read(), format='audio/wav')
            st.download_button("💾 Télécharger l'audio", open(temp_audio_path, "rb").read(), file_name="enregistrement.wav")

        except Exception as e:
            st.error(f"❌ Erreur lors de la création/écriture du fichier audio temporaire : {str(e)}")
            st.session_state.audio_file_path = None # S'assurer que l'état est propre
            st.session_state.audio_transcript = ""

def upload_audio_file():
    st.subheader("🎧 Uploader un fichier audio")
    uploaded_file = st.file_uploader(
        "Choisir un fichier audio",
        type=["mp3", "wav", "ogg", "flac", "aac", "m4a"],
        help="Formats acceptés : MP3, WAV, OGG, FLAC, AAC, M4A",
        key="audio_uploader"
    )

    if uploaded_file is not None:
        st.success("✅ Fichier audio uploadé !")
        
        # Utiliser tempfile pour créer un fichier temporaire pour l'audio uploadé
        try:
            st.info("Création du fichier temporaire...")
            # Utiliser le suffixe basé sur l'extension du fichier uploadé
            suffix = os.path.splitext(uploaded_file.name)[1]
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_audio_file:
                temp_audio_path = temp_audio_file.name
                st.info(f"Fichier temporaire créé : {temp_audio_path}")

                st.info("Écriture des données audio dans le fichier temporaire...")
                # uploaded_file.getvalue() contient les bytes du fichier uploadé
                temp_audio_file.write(uploaded_file.getvalue())
                st.info("Écriture terminée.")
            
            # Stocker le chemin du fichier temporaire dans la session
            st.session_state.audio_file_path = temp_audio_path
            st.session_state.audio_transcript = ""  # Reset transcript
            st.write(f"📂 Chemin local du fichier temporaire : `{temp_audio_path}`")

            # Afficher les options de lecture et téléchargement à partir du fichier temporaire
            st.audio(open(temp_audio_path, "rb").read(), format=uploaded_file.type) # Use uploaded file type
            st.download_button("💾 Télécharger l'audio", open(temp_audio_path, "rb").read(), file_name=uploaded_file.name)

        except Exception as e:
            st.error(f"❌ Erreur lors de la création/écriture du fichier audio temporaire : {str(e)}")
            st.session_state.audio_file_path = None # S'assurer que l'état est propre
            st.session_state.audio_transcript = ""

def parse_pv_text(pv_text, meeting_info):
    """
    Parse le texte du PV généré par Gemini pour extraire :
    - Les participants (Monsieur/Madame + nom)
    - Les points d'ordre du jour (titre + contenu)
    Retourne un dictionnaire structuré pour create_word_pv.
    """
    pv_struct = {
        'presents': [],
        'absents': [],
        'assistent': [],
        'ordre_du_jour': []
    }
    # Participants : on prend ceux de meeting_info (plus fiable)
    for name, title in meeting_info.get('participants', []):
        civilite = 'Monsieur' if 'M.' in title or 'Monsieur' in title else 'Madame' if 'Mme' in title or 'Madame' in title else 'Monsieur'
        pv_struct['presents'].append((civilite, name))
    # TODO : gérer absents et assistent si tu ajoutes ces champs dans l'UI

    # Points d'ordre du jour
    # On cherche les titres de points (ex: 1. TITRE)
    point_pattern = re.compile(r'\n?(\d+)\.\s*(.+?)(?=\n\d+\.|\Z)', re.DOTALL)
    matches = list(point_pattern.finditer(pv_text))
    for i, match in enumerate(matches):
        titre = match.group(2).strip().split('\n')[0]
        # Le contenu est tout ce qui suit le titre jusqu'au prochain point
        start = match.end()
        end = matches[i+1].start() if i+1 < len(matches) else len(pv_text)
        contenu = pv_text[start:end].strip()
        pv_struct['ordre_du_jour'].append({'titre': titre, 'contenu': contenu})
    return pv_struct

def main():
    st.set_page_config(
        page_title="Assistant de Réunion CMR",
        page_icon="📊",
        layout="wide"
    )

    # Configuration de l'API Gemini
    configure_apis()
    
    # Ajout du titre avec logo
    col1, col2, col3 = st.columns([2, 2, 8])
    with col2:
        st.image("logo.png", width=200)
    with col3:
        st.markdown(
            "<h1 style='text-align: left; margin-top: 25px;'>Génération du PV de Réunion</h1>",
            unsafe_allow_html=True
    )
    
    # Variables de session pour stocker les résultats
    if 'video_transcript' not in st.session_state:
        st.session_state.video_transcript = ""
    if 'handwritten_text' not in st.session_state:
        st.session_state.handwritten_text = ""
    if 'pdf_summary' not in st.session_state:
        st.session_state.pdf_summary = ""
    if 'pdf_data' not in st.session_state:
        st.session_state.pdf_data = {}
    if 'audio_transcript' not in st.session_state:
        st.session_state.audio_transcript = ""
    if 'meeting_info' not in st.session_state:
        st.session_state.meeting_info = None
    if 'additional_participants' not in st.session_state:
        st.session_state.additional_participants = []

    # Section d'informations de base du PV
    st.header("📝 Informations de base du PV")
    
    col1, col2 = st.columns(2)
    
    with col1:
        pv_number = st.text_input("Numéro du PV", "02/24")
        date = st.date_input("Date", format="DD/MM/YYYY")
        lieu = st.text_input("Lieu", " ")
        heure_debut = st.time_input("Heure début")
        heure_fin = st.time_input("Heure fin")
    
    with col2:
        st.subheader("Participants")
        participants = []
        
        # Option pour ajouter des participants supplémentaires
        if st.button("Ajouter un participant"):
            st.session_state.additional_participants.append(len(st.session_state.additional_participants))
        
        # Afficher les champs pour les participants supplémentaires
        for i in st.session_state.additional_participants:
            st.write("**Participant supplémentaire**")
            col1, col2, col3 = st.columns(3)
            with col1:
                name = st.text_input(f"Nom {i+1}")
            with col2:
                title = st.text_input(f"Titre {i+1}")
            if name and title:
                participants.append((name, title))

                

    # Stocker les informations de la réunion
    st.session_state.meeting_info = {
        'pv_number': pv_number,
        'date': date.strftime("%d/%m/%Y"),
        'lieu': lieu,
        'heure_debut': heure_debut.strftime("%H:%M"),
        'heure_fin': heure_fin.strftime("%H:%M"),
        'participants': [(name, title) for name, title in participants if name and title],
        'participants_by_section': {
            "PRÉSENTS OU REPRÉSENTÉS": [f"{name} - {title}" for name, title in participants if name and title],
            "ABSENTS EXCUSÉS": [],
            "ASSISTENT ÉGALEMENT": []
        }
    }

    # Section d'upload des fichiers
    st.markdown("### 📁 Importation des documents")

    audio_input_mode = st.radio(
        "Source audio :",
        ("Enregistrer l'audio", "Uploader un fichier audio"),
        horizontal=True,
        key="audio_input_mode"
    )

    if audio_input_mode == "Enregistrer l'audio":
        record_audio_simple()
    elif audio_input_mode == "Uploader un fichier audio":
        upload_audio_file()


    # Style CSS pour contrôler individuellement chaque drag and drop
    st.markdown("""
        <style>
        /* Styles de base pour tous les uploaders */
        .stFileUploader > div {
            display: flex;
            align-items: center;
            justify-content: center;
        }
        
        /* Style spécifique pour l'uploader vidéo */
        [data-testid="stFileUploader"]:has(#video_uploader) {
            height: 150px;
            margin-top: 0.5rem;
            margin-bottom: 1rem;
            background-color: rgba(255, 255, 255, 0.05);
        }
        
        /* Style spécifique pour l'uploader d'images */
        [data-testid="stFileUploader"]:has(#image_uploader) {
            height: 180px;
            margin-top: 0.5rem;
            margin-bottom: 1rem;
            background-color: rgba(255, 255, 255, 0.05);
        }
        
        /* Style spécifique pour l'uploader PDF */
        [data-testid="stFileUploader"]:has(#pdf_uploader) {
            height: 160px;
            margin-top: 0.75rem;
            margin-bottom: 1rem;
            background-color: rgba(255, 255, 255, 0.05);
        }
        
        /* Style pour le conteneur des colonnes */
        .row-widget.stHorizontalBlock {
            align-items: flex-start;
            gap: 1.5rem;
        }
        
        /* Style pour les titres des sections */
        .element-container h3 {
            margin-bottom: 0.75rem;
        }
        </style>
    """, unsafe_allow_html=True)
    


    # Créer d'abord les titres dans une rangée
    title_cols = st.columns(4)
    with title_cols[0]:
        st.markdown("### 🎥 Vidéo de la réunion")
    with title_cols[1]:
        st.markdown("### 📝 Images manuscrites")
    with title_cols[2]:
        st.markdown("### 📄 Documents PDF")
    
    # Ensuite, créer les options radio pour la vidéo dans une rangée séparée
    radio_col, empty_col1, empty_col2 = st.columns(3)
    with radio_col:
        video_upload_mode = st.radio(
            "Mode d'importation :",("Uploader un fichier", "Fournir un lien"),
            horizontal=True,
            key="video_mode"
        )

    # Ensuite, créer les textes d'instructions dans une rangée séparée
    text_cols = st.columns(3)
    with text_cols[0]:
        if video_upload_mode == "Uploader un fichier":
            st.markdown("Importez votre vidéo")
    with text_cols[1]:
        st.markdown("Importez vos images")
    with text_cols[2]:
        st.markdown("Importez vos documents")

    # Initialiser video_file et video_url à None pour éviter UnboundLocalError
    video_file = None
    video_url = None

    # Enfin, créer les zones de téléchargement dans une rangée séparée
    upload_cols = st.columns(3)
    with upload_cols[0]:
        if video_upload_mode == "Uploader un fichier":
            video_file = st.file_uploader(
                "Importer une vidéo",
                type=["mp4", "vro", "mpeg4"],
                help="Formats acceptés : MP4, VRO, MPEG4 • Limite : 2GB",
                key="video_uploader",
                label_visibility="collapsed"
            )
        else:
            video_url = st.text_input(
                "Lien de la vidéo",
                placeholder="https://drive.google.com/file/d/...",
                help="Lien Google Drive partagé"
            )
    with upload_cols[1]:
        image_files = st.file_uploader(
            "Importer des images",
            type=["jpg", "jpeg", "png"],
            accept_multiple_files=True,
            help="Formats acceptés : JPG, JPEG, PNG • Limite : 2GB par fichier",
            key="image_uploader",
            label_visibility="collapsed"
        )
    with upload_cols[2]:
        pdf_files = st.file_uploader(
            "Importer des PDF",
            type=["pdf"],
            accept_multiple_files=True,
            help="Format accepté : PDF • Limite : 2GB par fichier",
            key="pdf_uploader",
            label_visibility="collapsed"
        )
    

    # Bouton de démarrage centré avec espace au-dessus
    st.markdown("<div style='text-align: center; margin-top: 2em;'>", unsafe_allow_html=True)
    if st.button("🚀 Démarrer le traitement", use_container_width=True):
        if not st.session_state.meeting_info:
            st.error("❌ Veuillez remplir les informations de base du PV avant de commencer le traitement.")
            return
        # Créer des conteneurs pour les résultats
        video_container = st.container()
        images_container = st.container()
        pdfs_container = st.container()
        audio_container = st.container()
        pv_container = st.container()

        # Afficher spinner global
        global_status = st.info("⏳ Traitement en cours, veuillez patienter...")

        # Traitement de l'audio (transcription batch au moment du traitement)
        if hasattr(st.session_state, 'audio_file_path') and st.session_state.audio_file_path:
            audio_file_to_process = st.session_state.audio_file_path
            try:
                with audio_container:
                    st.subheader("🎤 Traitement de l'audio")
                    # Lire le fichier local pour l'affichage
                    st.audio(open(audio_file_to_process, "rb").read(), format='audio/wav')
                    status = st.info("Transcription de l'audio en cours...")
                    progress_bar = st.progress(0)

                    # Le traitement FFmpeg et segmentation se fera depuis le fichier WAV local
                    # Pas besoin de TemporaryDirectory ici, on utilise directement le fichier local
                    segments = segment_audio(audio_file_to_process) # segment_audio prend le chemin en paramètre

                    transcript = []
                    total = len(segments)
                    for i, segment_path in enumerate(segments):
                        # process_segment_batch prend des chemins de segments temporaires créés par segment_audio
                        # Note: segment_audio crée déjà des fichiers temporaires, process_segment_batch les lit et les supprime.
                        # La logique ici reste similaire, on passe les chemins des segments.
                        batch_result = process_segment_batch([segment_path], 0, 1, 1, progress_bar, status)
                        transcript.extend(batch_result)
                        # progress_bar et status sont mis à jour dans process_segment_batch maintenant

                    st.session_state.audio_transcript = "\n".join(transcript)
                    status.success("✅ Transcription audio terminée!")
                    st.text_area("Transcription de l'audio:", st.session_state.audio_transcript, height=200)

            except Exception as e:
                st.error(f"❌ Erreur lors du traitement de l'audio : {str(e)}")
                # S'assurer que l'état de la transcription est vide en cas d'erreur
                st.session_state.audio_transcript = ""
            finally:
                # Nettoyer le fichier WAV local après traitement (réussi ou non)
                if os.path.exists(audio_file_to_process):
                    try:
                        os.remove(audio_file_to_process)
                        st.info(f"Fichier audio local supprimé : {audio_file_to_process}")
                    except Exception as e:
                        st.warning(f"⚠️ Impossible de supprimer le fichier audio local {audio_file_to_process}: {str(e)}")
                # Réinitialiser le chemin dans la session
                st.session_state.audio_file_path = None

        # Traitement de la vidéo
        with video_container:
            st.subheader("🎥 Traitement de la vidéo")
            if (video_file is not None) or (video_url is not None and video_url.strip() != ""):
          # Vérifier si on a soit un fichier soit une URL valide
                with st.spinner("Transcription en cours..."):
                    if video_file:
                        st.session_state.video_transcript = transcribe_video(video_file)
                    elif video_url and video_url.strip():
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as temp_video:
                            temp_video_path = temp_video.name
                            if download_video_from_drive(video_url, temp_video_path):
                                if verify_video_file(temp_video_path):
                                    st.session_state.video_transcript = transcribe_video(open(temp_video_path, "rb"))
                                else:
                                    st.error("❌ Le fichier vidéo téléchargé n'est pas valide")
                            else:
                                st.error("❌ Échec du téléchargement de la vidéo")

                    if st.session_state.video_transcript:
                        st.success("✅ Transcription terminée!")
                        st.text_area("Transcription:", st.session_state.video_transcript, height=200)
            else:
                st.info("ℹ️ Aucune vidéo n'a été fournie")

        # Traitement des images
        if image_files:
            with images_container:
                st.subheader("🖼️ Traitement des images")
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                transcriptions = {}
                for idx, image_file in enumerate(image_files):
                    try:
                        status_text.text(f"Analyse de l'image {idx + 1}/{len(image_files)}: {image_file.name}")
                        image_bytes = image_file.read()
                        image = Image.open(io.BytesIO(image_bytes))
                        
                        col1, col2 = st.columns([1, 1])
                        with col1:
                            st.image(image, caption=f"Image {idx + 1}: {image_file.name}", use_column_width=True)
                        with col2:
                            transcription = process_handwritten_image(image_bytes)
                            if transcription:
                                st.text_area(f"Texte reconnu - Image {idx + 1}", transcription, height=150)
                                transcriptions[image_file.name] = transcription
                        
                        progress_bar.progress((idx + 1)/len(image_files))
                    except Exception as e:
                        st.error(f"❌ Erreur lors du traitement de l'image {image_file.name}: {str(e)}")
                
                if transcriptions:
                    st.session_state.handwritten_text = "\n\n".join([f"[Image: {name}]\n{text}" for name, text in transcriptions.items()])
                    st.success("✅ Traitement des images terminé!")

        # Traitement des PDFs
        if pdf_files:
            with pdfs_container:
                st.subheader("📄 Traitement des PDFs")
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                # Réinitialiser les données PDF de la session
                st.session_state.pdf_data = {}
                pdf_summaries_list = [] # Liste temporaire pour l'ancien format
                
                for idx, pdf_file in enumerate(pdf_files):
                    try:
                        status_text.text(f"Analyse du document {idx + 1}/{len(pdf_files)}: {pdf_file.name}")
                        # Lire à nouveau car le pointeur peut être à la fin après l'upload
                        pdf_file.seek(0) 
                        pdf_result = process_pdf(pdf_file)
                        
                        if pdf_result["summary"]:
                            # Stocker le résultat structuré
                            st.session_state.pdf_data[pdf_file.name] = pdf_result
                            # Ajouter au résumé global pour generate_meeting_minutes
                            pdf_summaries_list.append(f"[Document: {pdf_file.name}]\n{pdf_result['summary']}")
                            
                            # Afficher l'aperçu
                            with st.expander(f"📄 Document {idx + 1}: {pdf_file.name} (Analysé)"):
                                st.text_area("Aperçu du contenu extrait:", pdf_result["summary"], height=200)
                                if pdf_result["acronyms"]:
                                    st.write("**Acronymes détectés:**")
                                    st.json(pdf_result["acronyms"])
                                else:
                                    st.write("Aucun acronyme détecté.")
                        else:
                             st.warning(f"Aucun contenu extrait pour {pdf_file.name}")
                        
                        progress_bar.progress((idx + 1)/len(pdf_files))
                    except Exception as e:
                        st.error(f"❌ Erreur lors de l'analyse du PDF {pdf_file.name}: {str(e)}")
                        # Stocker une indication d'erreur
                        st.session_state.pdf_data[pdf_file.name] = {"summary": f"[Erreur: {str(e)}]", "acronyms": {}} 
                        pdf_summaries_list.append(f"[Document: {pdf_file.name}]\n[Erreur lors de l'analyse: {str(e)}]")
                
                # Mettre à jour l'ancien état pdf_summary pour generate_meeting_minutes
                st.session_state.pdf_summary = "\n\n".join(pdf_summaries_list)
                if st.session_state.pdf_data:
                    st.success("✅ Traitement des PDFs terminé!")
                else:
                    st.warning("Aucun PDF n'a pu être traité.")

        # Génération du PV (uniquement après la fin de la transcription)
        with pv_container:
            st.subheader("Génération du PV")
            pdf_summary_for_generation = "\n\n".join(
                [f"[Document: {name}]\n{data.get('summary', '')}" 
                 for name, data in st.session_state.get('pdf_data', {}).items()]
            )
            if any([st.session_state.video_transcript, st.session_state.handwritten_text, pdf_summary_for_generation, st.session_state.get("audio_transcript", "")]):
                with st.spinner("Génération du PV en cours..."):
                    pv = generate_meeting_minutes(
                        st.session_state.video_transcript,
                        st.session_state.handwritten_text,
                        pdf_summary_for_generation,
                        st.session_state.meeting_info,
                        st.session_state.get("audio_transcript", "")
                    )
                    if pv:
                        st.success("✅ PV généré avec succès!")
                        st.text_area("Procès-verbal de la réunion:", pv, height=500)
                        try:
                            # Créer le document Word directement à partir du texte généré
                            doc_buffer = create_word_pv(pv, st.session_state.meeting_info)
                            st.download_button(
                                label="📎 Télécharger le PV en format Word",
                                data=doc_buffer,
                                file_name=f"PV_{st.session_state.meeting_info.get('pv_number', 'NA').replace('/', '_')}_Comite_Audit.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        except Exception as e_word:
                            st.error(f"❌ Erreur lors de la création du document Word: {str(e_word)}")
                global_status.success("✅ Traitement terminé !")
            else:
                global_status.warning("⚠️ Aucun contenu à traiter pour générer le PV")

if __name__ == "__main__":
    main() 
